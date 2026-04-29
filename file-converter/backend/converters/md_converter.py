"""
Markdown → PDF / DOCX / TXT / XLSX
"""

import io
import re


def convert_md(content: bytes, output_format: str) -> bytes:
    text = content.decode("utf-8", errors="replace")

    if output_format == "txt":
        return _md_to_txt(text)
    elif output_format == "pdf":
        return _md_to_pdf(text)
    elif output_format == "docx":
        return _md_to_docx(text)
    elif output_format == "xlsx":
        return _md_to_xlsx(text)
    else:
        raise ValueError(f"Unsupported output format: {output_format}")


# ── TXT ───────────────────────────────────────────────────────
def _md_to_txt(text: str) -> bytes:
    """Strip Markdown syntax to plain text."""
    # Remove code fences
    text = re.sub(r"```[\s\S]*?```", lambda m: m.group(0).replace("```", ""), text)
    # Remove inline code
    text = re.sub(r"`([^`]+)`", r"\1", text)
    # Remove images
    text = re.sub(r"!\[([^\]]*)\]\([^\)]*\)", r"\1", text)
    # Remove links but keep text
    text = re.sub(r"\[([^\]]+)\]\([^\)]+\)", r"\1", text)
    # Remove bold/italic
    text = re.sub(r"\*{1,3}([^*]+)\*{1,3}", r"\1", text)
    text = re.sub(r"_{1,3}([^_]+)_{1,3}", r"\1", text)
    # Remove heading markers
    text = re.sub(r"^#{1,6}\s+", "", text, flags=re.MULTILINE)
    # Remove horizontal rules
    text = re.sub(r"^[-*_]{3,}\s*$", "", text, flags=re.MULTILINE)
    # Remove blockquote markers
    text = re.sub(r"^>\s?", "", text, flags=re.MULTILINE)
    # Clean up extra blank lines
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip().encode("utf-8")


# ── PDF ───────────────────────────────────────────────────────
def _md_to_pdf(text: str) -> bytes:
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Preformatted
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.pagesizes import LETTER
    from reportlab.lib.units import inch
    from reportlab.lib import colors
    import markdown as mdlib

    buf = io.BytesIO()
    doc = SimpleDocTemplate(
        buf,
        pagesize=LETTER,
        leftMargin=1 * inch,
        rightMargin=1 * inch,
        topMargin=1 * inch,
        bottomMargin=1 * inch,
    )

    styles = getSampleStyleSheet()
    _add_custom_styles(styles)

    # Convert MD to HTML then parse into ReportLab flowables
    html = mdlib.markdown(text, extensions=["tables", "fenced_code"])
    story = _html_to_flowables(html, styles)

    doc.build(story)
    return buf.getvalue()


def _add_custom_styles(styles):
    from reportlab.lib.styles import ParagraphStyle
    from reportlab.lib import colors
    from reportlab.lib.enums import TA_LEFT

    for tag, size, bold in [
        ("h1", 22, True), ("h2", 18, True), ("h3", 15, True),
        ("h4", 13, True), ("h5", 11, True), ("h6", 10, True),
    ]:
        if tag not in styles:
            styles.add(ParagraphStyle(
                name=tag,
                parent=styles["Normal"],
                fontSize=size,
                leading=size * 1.3,
                fontName="Helvetica-Bold" if bold else "Helvetica",
                spaceAfter=6,
                spaceBefore=12,
            ))

    if "Code" not in styles:
        styles.add(ParagraphStyle(
            name="Code",
            parent=styles["Normal"],
            fontName="Courier",
            fontSize=9,
            backColor=colors.HexColor("#f4f4f4"),
            borderPadding=(4, 4, 4, 4),
        ))


def _html_to_flowables(html: str, styles):
    """Very small HTML→flowable converter sufficient for typical Markdown output."""
    from reportlab.platypus import Paragraph, Spacer, Preformatted
    from reportlab.lib.units import inch
    from html.parser import HTMLParser

    class _P(HTMLParser):
        def __init__(self):
            super().__init__()
            self.flowables = []
            self._buf = []
            self._tag_stack = []
            self._in_pre = False

        def handle_starttag(self, tag, attrs):
            self._tag_stack.append(tag)
            if tag == "pre":
                self._in_pre = True
            elif tag in ("h1", "h2", "h3", "h4", "h5", "h6", "p", "li"):
                self._buf = []
            elif tag == "br":
                self._buf.append("<br/>")

        def handle_endtag(self, tag):
            if self._tag_stack and self._tag_stack[-1] == tag:
                self._tag_stack.pop()

            if tag in ("h1", "h2", "h3", "h4", "h5", "h6"):
                text = "".join(self._buf).strip()
                if text:
                    style_name = tag if tag in styles else "Heading1"
                    self.flowables.append(Paragraph(text, styles.get(tag, styles["Heading1"])))
                    self.flowables.append(Spacer(1, 0.1 * inch))
            elif tag in ("p",):
                text = "".join(self._buf).strip()
                if text:
                    self.flowables.append(Paragraph(text, styles["Normal"]))
                    self.flowables.append(Spacer(1, 0.05 * inch))
            elif tag in ("li",):
                text = "• " + "".join(self._buf).strip()
                if text.strip():
                    self.flowables.append(Paragraph(text, styles["Normal"]))
            elif tag == "pre":
                self._in_pre = False
                text = "".join(self._buf).strip()
                self.flowables.append(Preformatted(text, styles.get("Code", styles["Code"])))
                self.flowables.append(Spacer(1, 0.05 * inch))

        def handle_data(self, data):
            if self._tag_stack and self._tag_stack[-1] in ("script", "style"):
                return
            self._buf.append(data)

    p = _P()
    p.feed(html)
    return p.flowables or [Paragraph("(empty)", styles["Normal"])]


# ── DOCX ──────────────────────────────────────────────────────
def _md_to_docx(text: str) -> bytes:
    from docx import Document
    from docx.shared import Pt
    import markdown as mdlib
    from html.parser import HTMLParser

    doc = Document()
    html = mdlib.markdown(text, extensions=["tables", "fenced_code"])

    class _P(HTMLParser):
        def __init__(self, doc):
            super().__init__()
            self.doc = doc
            self._buf = []
            self._tag = None

        def handle_starttag(self, tag, attrs):
            self._tag = tag
            self._buf = []

        def handle_endtag(self, tag):
            text = "".join(self._buf).strip()
            if not text:
                return
            if tag in ("h1", "h2", "h3"):
                lvl = int(tag[1])
                self.doc.add_heading(text, level=lvl)
            elif tag == "p":
                self.doc.add_paragraph(text)
            elif tag == "li":
                self.doc.add_paragraph(text, style="List Bullet")
            elif tag in ("pre", "code"):
                p = self.doc.add_paragraph(text)
                p.runs[0].font.name = "Courier New"
                p.runs[0].font.size = Pt(9)
            self._tag = None

        def handle_data(self, data):
            self._buf.append(data)

    _P(doc).feed(html)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ── XLSX ──────────────────────────────────────────────────────
def _md_to_xlsx(text: str) -> bytes:
    """Extract tables from Markdown; fall back to line-by-line if none found."""
    import openpyxl
    from openpyxl.styles import Font, PatternFill

    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "Markdown"

    lines = text.splitlines()
    row = 1

    # Try to detect and parse MD tables
    table_block = []
    in_table = False

    for line in lines:
        stripped = line.strip()
        if "|" in stripped:
            in_table = True
            table_block.append(stripped)
        else:
            if in_table and table_block:
                _write_md_table(ws, table_block, row)
                row += len(table_block) + 1
                table_block = []
                in_table = False
            # Write non-table lines
            ws.cell(row=row, column=1, value=stripped or None)
            row += 1

    if table_block:
        _write_md_table(ws, table_block, row)

    # Auto-size columns
    for col in ws.columns:
        max_len = max((len(str(c.value or "")) for c in col), default=10)
        ws.column_dimensions[col[0].column_letter].width = min(max_len + 2, 60)

    buf = io.BytesIO()
    wb.save(buf)
    return buf.getvalue()


def _write_md_table(ws, lines, start_row):
    from openpyxl.styles import Font, PatternFill, Alignment

    header_fill = PatternFill("solid", fgColor="1a1a2e")
    header_font = Font(color="c9a84c", bold=True)

    data_rows = [l for l in lines if not re.match(r"^\|[-:| ]+\|$", l)]

    for i, line in enumerate(data_rows):
        cols = [c.strip() for c in line.strip().strip("|").split("|")]
        for j, val in enumerate(cols, start=1):
            cell = ws.cell(row=start_row + i, column=j, value=val)
            if i == 0:
                cell.fill = header_fill
                cell.font = header_font
            cell.alignment = Alignment(wrap_text=True)
