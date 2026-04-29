"""
JSON → PDF / DOCX / TXT / XLSX
"""

import io
import json


def convert_json(content: bytes, output_format: str) -> bytes:
    text = content.decode("utf-8", errors="replace")
    try:
        data = json.loads(text)
    except json.JSONDecodeError as e:
        raise ValueError(f"Invalid JSON: {e}")

    if output_format == "txt":
        return _json_to_txt(data)
    elif output_format == "pdf":
        return _json_to_pdf(data)
    elif output_format == "docx":
        return _json_to_docx(data)
    elif output_format == "xlsx":
        return _json_to_xlsx(data)
    else:
        raise ValueError(f"Unsupported output format: {output_format}")


# ── TXT ───────────────────────────────────────────────────────
def _json_to_txt(data) -> bytes:
    return json.dumps(data, indent=2, ensure_ascii=False).encode("utf-8")


# ── PDF ───────────────────────────────────────────────────────
def _json_to_pdf(data) -> bytes:
    from reportlab.platypus import SimpleDocTemplate, Preformatted, Paragraph, Spacer
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.pagesizes import LETTER
    from reportlab.lib.units import inch
    from reportlab.lib import colors

    buf = io.BytesIO()
    doc = SimpleDocTemplate(
        buf, pagesize=LETTER,
        leftMargin=inch, rightMargin=inch,
        topMargin=inch, bottomMargin=inch,
    )
    styles = getSampleStyleSheet()

    code_style = ParagraphStyle(
        "JSONCode",
        parent=styles["Normal"],
        fontName="Courier",
        fontSize=8.5,
        leading=12,
        backColor=colors.HexColor("#f8f8f8"),
        borderColor=colors.HexColor("#cccccc"),
        borderWidth=0.5,
        borderPadding=6,
        wordWrap="LTR",
    )

    pretty = json.dumps(data, indent=2, ensure_ascii=False)

    story = []

    # Summary header
    dtype = _describe_type(data)
    story.append(Paragraph(f"<b>JSON Document</b> — {dtype}", styles["Heading2"]))
    story.append(Spacer(1, 0.2 * inch))

    # If root is array of objects, add summary table
    if isinstance(data, list) and data and isinstance(data[0], dict):
        story.extend(_table_flowables(data[:50], styles))
        story.append(Spacer(1, 0.2 * inch))

    # Full JSON dump
    story.append(Paragraph("<b>Full JSON</b>", styles["Heading3"]))
    story.append(Spacer(1, 0.1 * inch))

    # Split large JSON across paragraphs (avoids single massive Preformatted overflow)
    lines = pretty.splitlines()
    chunk_size = 40
    for i in range(0, len(lines), chunk_size):
        chunk = "\n".join(lines[i:i + chunk_size])
        story.append(Preformatted(chunk, code_style))

    doc.build(story)
    return buf.getvalue()


def _describe_type(data) -> str:
    if isinstance(data, list):
        return f"Array of {len(data)} items"
    if isinstance(data, dict):
        return f"Object with {len(data)} keys"
    return type(data).__name__


def _table_flowables(rows: list, styles):
    from reportlab.platypus import Table, TableStyle, Paragraph, Spacer
    from reportlab.lib import colors
    from reportlab.lib.units import inch

    all_keys = list(dict.fromkeys(k for row in rows for k in (row if isinstance(row, dict) else {})))
    if not all_keys:
        return []

    header = [Paragraph(f"<b>{k}</b>", styles["Normal"]) for k in all_keys]
    table_data = [header]

    for row in rows:
        if isinstance(row, dict):
            table_data.append([
                Paragraph(str(row.get(k, ""))[:120], styles["Normal"])
                for k in all_keys
            ])

    if len(table_data) < 2:
        return []

    col_w = [1.5 * inch] * min(len(all_keys), 5)
    t = Table(table_data[:51], colWidths=col_w, repeatRows=1)
    t.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#1a1a2e")),
        ("TEXTCOLOR", (0, 0), (-1, 0), colors.HexColor("#c9a84c")),
        ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
        ("FONTSIZE", (0, 0), (-1, -1), 8),
        ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, colors.HexColor("#f8f8f8")]),
        ("GRID", (0, 0), (-1, -1), 0.4, colors.HexColor("#cccccc")),
        ("VALIGN", (0, 0), (-1, -1), "TOP"),
        ("LEFTPADDING", (0, 0), (-1, -1), 4),
        ("RIGHTPADDING", (0, 0), (-1, -1), 4),
    ]))

    return [
        Paragraph("<b>Data Preview (first 50 rows)</b>", styles["Heading3"]),
        Spacer(1, 0.1 * inch),
        t,
    ]


# ── DOCX ──────────────────────────────────────────────────────
def _json_to_docx(data) -> bytes:
    from docx import Document
    from docx.shared import Pt, RGBColor

    doc = Document()
    doc.add_heading("JSON Document", 1)

    if isinstance(data, list) and data and isinstance(data[0], dict):
        keys = list(dict.fromkeys(k for row in data for k in (row if isinstance(row, dict) else {})))
        table = doc.add_table(rows=1, cols=len(keys))
        table.style = "Light List Accent 1"

        for j, k in enumerate(keys):
            table.rows[0].cells[j].text = str(k)

        for row in data[:200]:
            cells = table.add_row().cells
            for j, k in enumerate(keys):
                cells[j].text = str(row.get(k, ""))

        doc.add_paragraph()
    elif isinstance(data, dict):
        table = doc.add_table(rows=1, cols=2)
        table.style = "Light List Accent 1"
        table.rows[0].cells[0].text = "Key"
        table.rows[0].cells[1].text = "Value"
        for k, v in data.items():
            row = table.add_row()
            row.cells[0].text = str(k)
            row.cells[1].text = json.dumps(v) if isinstance(v, (dict, list)) else str(v)
        doc.add_paragraph()

    # Raw JSON
    doc.add_heading("Raw JSON", 2)
    p = doc.add_paragraph(json.dumps(data, indent=2, ensure_ascii=False))
    p.runs[0].font.name = "Courier New"
    p.runs[0].font.size = Pt(8)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ── XLSX ──────────────────────────────────────────────────────
def _json_to_xlsx(data) -> bytes:
    import openpyxl
    from openpyxl.styles import Font, PatternFill, Alignment, Border, Side

    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "Data"

    header_fill = PatternFill("solid", fgColor="1a1a2e")
    header_font = Font(color="c9a84c", bold=True)
    thin = Side(style="thin", color="cccccc")
    border = Border(left=thin, right=thin, top=thin, bottom=thin)

    if isinstance(data, list) and data:
        if isinstance(data[0], dict):
            keys = list(dict.fromkeys(k for row in data for k in (row if isinstance(row, dict) else {})))

            for j, k in enumerate(keys, start=1):
                cell = ws.cell(row=1, column=j, value=str(k))
                cell.fill = header_fill
                cell.font = header_font
                cell.border = border

            for i, row in enumerate(data, start=2):
                if isinstance(row, dict):
                    for j, k in enumerate(keys, start=1):
                        val = row.get(k, "")
                        if isinstance(val, (dict, list)):
                            val = json.dumps(val)
                        cell = ws.cell(row=i, column=j, value=val)
                        cell.border = border
                        cell.alignment = Alignment(wrap_text=True)
        else:
            ws.cell(row=1, column=1, value="Value").font = header_font
            for i, v in enumerate(data, start=2):
                ws.cell(row=i, column=1, value=json.dumps(v) if isinstance(v, (dict, list)) else v)

    elif isinstance(data, dict):
        ws.cell(row=1, column=1, value="Key").fill = header_fill
        ws.cell(row=1, column=1).font = header_font
        ws.cell(row=1, column=2, value="Value").fill = header_fill
        ws.cell(row=1, column=2).font = header_font

        for i, (k, v) in enumerate(data.items(), start=2):
            ws.cell(row=i, column=1, value=str(k))
            ws.cell(row=i, column=2, value=json.dumps(v) if isinstance(v, (dict, list)) else str(v))
    else:
        ws.cell(row=1, column=1, value=str(data))

    # Auto-size
    for col in ws.columns:
        max_len = max((len(str(c.value or "")) for c in col), default=10)
        ws.column_dimensions[col[0].column_letter].width = min(max_len + 2, 50)

    buf = io.BytesIO()
    wb.save(buf)
    return buf.getvalue()
