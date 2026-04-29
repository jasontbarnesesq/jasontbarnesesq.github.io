"""
BAS (BASIC source) / TXT → PDF / DOCX / TXT / XLSX

BAS files are treated as plain-text source code.
PDF and DOCX render with syntax highlighting (keyword coloring).
XLSX exports each line as a row with optional token analysis.
"""

import io
import re


# Keywords common across BASIC dialects
BASIC_KEYWORDS = {
    "statement": {
        "PRINT", "LET", "INPUT", "IF", "THEN", "ELSE", "ELSEIF", "END",
        "FOR", "TO", "STEP", "NEXT", "WHILE", "WEND", "DO", "LOOP",
        "UNTIL", "GOTO", "GOSUB", "RETURN", "ON", "DIM", "REDIM",
        "OPEN", "CLOSE", "READ", "WRITE", "DATA", "RESTORE", "STOP",
        "REM", "SUB", "FUNCTION", "CALL", "DEF", "EXIT", "SELECT",
        "CASE", "WITH", "EACH", "IN", "CLASS", "MODULE", "TYPE",
        "PUBLIC", "PRIVATE", "STATIC", "AS", "NEW", "SET", "GET",
        "PROPERTY", "BYREF", "BYVAL", "OPTIONAL", "PARAMARRAY",
    },
    "builtin": {
        "ABS", "ATN", "CHR", "COS", "EXP", "FIX", "INT", "LEN",
        "LOG", "MID", "SIN", "SQR", "STR", "TAN", "VAL", "LEFT",
        "RIGHT", "LTRIM", "RTRIM", "TRIM", "UCASE", "LCASE",
        "INSTR", "SPACE", "STRING", "DATE", "TIME", "NOW",
        "ISNULL", "ISEMPTY", "ISARRAY", "ISNUMERIC",
    },
}

ALL_KEYWORDS = {k for s in BASIC_KEYWORDS.values() for k in s}


def convert_bas(content: bytes, output_format: str) -> bytes:
    text = content.decode("utf-8", errors="replace")

    if output_format == "txt":
        return text.encode("utf-8")
    elif output_format == "pdf":
        return _bas_to_pdf(text)
    elif output_format == "docx":
        return _bas_to_docx(text)
    elif output_format == "xlsx":
        return _bas_to_xlsx(text)
    else:
        raise ValueError(f"Unsupported output format: {output_format}")


# ── Helpers ───────────────────────────────────────────────────
def _tokenize_line(line: str):
    """
    Returns list of (token_text, token_type) where token_type is
    'keyword', 'builtin', 'string', 'comment', 'number', 'operator', or 'plain'.
    """
    tokens = []
    i = 0
    while i < len(line):
        # Comments (REM or ')
        rem_match = re.match(r"(REM\b.*|'.*)", line[i:], re.IGNORECASE)
        if rem_match:
            tokens.append((rem_match.group(0), "comment"))
            break

        # String literals
        if line[i] == '"':
            j = line.find('"', i + 1)
            end = j + 1 if j != -1 else len(line)
            tokens.append((line[i:end], "string"))
            i = end
            continue

        # Numbers
        num_match = re.match(r"\d+(\.\d*)?([eE][+-]?\d+)?", line[i:])
        if num_match and num_match.group(0):
            tokens.append((num_match.group(0), "number"))
            i += len(num_match.group(0))
            continue

        # Identifiers / keywords
        id_match = re.match(r"[A-Za-z_]\w*\$?", line[i:])
        if id_match:
            word = id_match.group(0)
            upper = word.upper().rstrip("$")
            if upper in BASIC_KEYWORDS["statement"]:
                tokens.append((word, "keyword"))
            elif upper in BASIC_KEYWORDS["builtin"]:
                tokens.append((word, "builtin"))
            else:
                tokens.append((word, "plain"))
            i += len(word)
            continue

        # Operators / punctuation
        op_match = re.match(r"[+\-*/\\^=<>(),:;#&!%@?]", line[i:])
        if op_match:
            tokens.append((op_match.group(0), "operator"))
            i += 1
            continue

        tokens.append((line[i], "plain"))
        i += 1

    return tokens


# ── PDF ───────────────────────────────────────────────────────
_PDF_COLORS = {
    "keyword":  "#569cd6",  # blue
    "builtin":  "#4ec9b0",  # teal
    "string":   "#ce9178",  # orange-red
    "comment":  "#6a9955",  # green
    "number":   "#b5cea8",  # light green
    "operator": "#d4d4d4",  # light grey
    "plain":    "#d4d4d4",
}

_BG_DARK = "#1e1e1e"
_FG_DEFAULT = "#d4d4d4"


def _bas_to_pdf(text: str) -> bytes:
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, HRFlowable
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.pagesizes import LETTER
    from reportlab.lib.units import inch
    from reportlab.lib import colors

    buf = io.BytesIO()
    doc = SimpleDocTemplate(
        buf, pagesize=LETTER,
        leftMargin=0.75 * inch, rightMargin=0.75 * inch,
        topMargin=inch, bottomMargin=inch,
    )
    styles = getSampleStyleSheet()

    title_style = ParagraphStyle(
        "BASTitle", parent=styles["Heading1"],
        fontName="Helvetica-Bold", fontSize=16,
        textColor=colors.HexColor("#c9a84c"), spaceBefore=0, spaceAfter=8,
    )
    line_style = ParagraphStyle(
        "BASLine", parent=styles["Normal"],
        fontName="Courier", fontSize=8.5, leading=13,
        textColor=colors.HexColor(_FG_DEFAULT),
        backColor=colors.HexColor(_BG_DARK),
        leftIndent=4, rightIndent=4,
        spaceBefore=0, spaceAfter=0,
    )

    story = [
        Paragraph("BASIC Source Listing", title_style),
        HRFlowable(width="100%", thickness=1, color=colors.HexColor("#c9a84c"), spaceAfter=8),
    ]

    for lineno, raw_line in enumerate(text.splitlines(), start=1):
        tokens = _tokenize_line(raw_line)
        parts = [f'<font color="#555555">{lineno:4d} </font>']
        for tok_text, tok_type in tokens:
            color = _PDF_COLORS.get(tok_type, _FG_DEFAULT)
            escaped = (tok_text
                       .replace("&", "&amp;")
                       .replace("<", "&lt;")
                       .replace(">", "&gt;"))
            parts.append(f'<font color="{color}">{escaped}</font>')

        html_line = "".join(parts) or "&nbsp;"
        story.append(Paragraph(html_line, line_style))

    doc.build(story)
    return buf.getvalue()


# ── DOCX ──────────────────────────────────────────────────────
def _bas_to_docx(text: str) -> bytes:
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    DOCX_COLORS = {
        "keyword":  RGBColor(0x56, 0x9c, 0xd6),
        "builtin":  RGBColor(0x4e, 0xc9, 0xb0),
        "string":   RGBColor(0xce, 0x91, 0x78),
        "comment":  RGBColor(0x6a, 0x99, 0x55),
        "number":   RGBColor(0xb5, 0xce, 0xa8),
        "operator": RGBColor(0xd4, 0xd4, 0xd4),
        "plain":    RGBColor(0xd4, 0xd4, 0xd4),
    }

    doc = Document()
    doc.add_heading("BASIC Source Listing", 1)

    for raw_line in text.splitlines():
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)

        for tok_text, tok_type in _tokenize_line(raw_line):
            run = p.add_run(tok_text)
            run.font.name = "Courier New"
            run.font.size = Pt(9)
            run.font.color.rgb = DOCX_COLORS.get(tok_type, DOCX_COLORS["plain"])

        if not raw_line:
            p.add_run(" ").font.name = "Courier New"

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ── XLSX ──────────────────────────────────────────────────────
def _bas_to_xlsx(text: str) -> bytes:
    import openpyxl
    from openpyxl.styles import Font, PatternFill, Alignment

    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "Source"

    # Headers
    hfont = Font(bold=True, color="c9a84c")
    hfill = PatternFill("solid", fgColor="1a1a2e")

    for col, header in enumerate(["Line", "Content", "Keywords Used", "Strings", "Comments"], 1):
        c = ws.cell(row=1, column=col, value=header)
        c.font = hfont
        c.fill = hfill

    ws.column_dimensions["A"].width = 6
    ws.column_dimensions["B"].width = 60
    ws.column_dimensions["C"].width = 30
    ws.column_dimensions["D"].width = 25
    ws.column_dimensions["E"].width = 35

    for lineno, raw_line in enumerate(text.splitlines(), start=1):
        tokens = _tokenize_line(raw_line)
        keywords_used = ", ".join(
            t for t, typ in tokens if typ in ("keyword", "builtin")
        )
        strings_found = " ".join(t for t, typ in tokens if typ == "string")
        comments_found = " ".join(t for t, typ in tokens if typ == "comment")

        row = lineno + 1
        ws.cell(row=row, column=1, value=lineno)
        ws.cell(row=row, column=2, value=raw_line).font = Font(name="Courier New", size=9)
        ws.cell(row=row, column=3, value=keywords_used or None)
        ws.cell(row=row, column=4, value=strings_found or None)
        ws.cell(row=row, column=5, value=comments_found or None)

        for col in range(1, 6):
            ws.cell(row=row, column=col).alignment = Alignment(wrap_text=False)

        if lineno % 2 == 0:
            for col in range(1, 6):
                ws.cell(row=row, column=col).fill = PatternFill("solid", fgColor="f0f0f0")

    buf = io.BytesIO()
    wb.save(buf)
    return buf.getvalue()
